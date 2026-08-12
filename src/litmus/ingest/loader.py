"""Multi-format document loading: txt, md, pdf, docx, html, csv.

Every loader returns plain text plus a doc_id derived from the filename
(stem, so ``pricing.md`` -> ``pricing``). Encoding for text-based formats is
sniffed with ``chardet`` before falling back to utf-8 with replacement.
"""

from __future__ import annotations

import csv
import io
from dataclasses import dataclass
from pathlib import Path

SUPPORTED_EXTENSIONS = {".txt", ".md", ".pdf", ".docx", ".html", ".htm", ".csv"}


@dataclass
class LoadedDocument:
    doc_id: str
    path: str
    text: str


def _read_text_with_encoding_detection(path: Path) -> str:
    raw = path.read_bytes()
    try:
        import chardet

        detected = chardet.detect(raw)
        encoding = detected.get("encoding") or "utf-8"
    except Exception:  # noqa: BLE001
        encoding = "utf-8"
    try:
        return raw.decode(encoding)
    except (UnicodeDecodeError, LookupError):
        return raw.decode("utf-8", errors="replace")


def _load_pdf(path: Path) -> str:
    import fitz

    doc = fitz.open(str(path))
    try:
        return "\n\n".join(page.get_text() for page in doc)
    finally:
        doc.close()


def _load_docx(path: Path) -> str:
    import docx

    document = docx.Document(str(path))
    parts = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n\n".join(parts)


def _load_html(path: Path) -> str:
    from bs4 import BeautifulSoup

    raw = _read_text_with_encoding_detection(path)
    soup = BeautifulSoup(raw, "html.parser")
    for tag in soup(["script", "style"]):
        tag.decompose()
    return soup.get_text(separator="\n")


def _load_csv(path: Path) -> str:
    raw = _read_text_with_encoding_detection(path)
    reader = csv.reader(io.StringIO(raw))
    rows = list(reader)
    if not rows:
        return ""
    header, *body = rows
    lines = []
    for row in body:
        pairs = [f"{h}: {v}" for h, v in zip(header, row)]
        lines.append("; ".join(pairs))
    return "\n".join(lines)


_LOADERS = {
    ".txt": _read_text_with_encoding_detection,
    ".md": _read_text_with_encoding_detection,
    ".pdf": _load_pdf,
    ".docx": _load_docx,
    ".html": _load_html,
    ".htm": _load_html,
    ".csv": _load_csv,
}


def load_document(path: str | Path) -> LoadedDocument:
    path = Path(path)
    ext = path.suffix.lower()
    if ext not in _LOADERS:
        raise ValueError(f"Unsupported document type: {ext} ({path})")
    text = _LOADERS[ext](path)
    return LoadedDocument(doc_id=path.stem, path=str(path), text=text)


def load_directory(docs_dir: str | Path) -> list[LoadedDocument]:
    docs_dir = Path(docs_dir)
    if not docs_dir.is_dir():
        raise ValueError(f"docs_dir is not a directory: {docs_dir}")
    docs = []
    for path in sorted(docs_dir.rglob("*")):
        if path.is_file() and path.suffix.lower() in SUPPORTED_EXTENSIONS:
            try:
                docs.append(load_document(path))
            except Exception as exc:  # noqa: BLE001
                raise ValueError(f"Failed to load {path}: {exc}") from exc
    if not docs:
        raise ValueError(f"No supported documents found in {docs_dir}")
    return docs
